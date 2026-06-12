from __future__ import annotations

from pathlib import Path

from pydantic import BaseModel, Field


class ExtractedDocument(BaseModel):
    source_type: str
    uri: str
    title: str
    text: str
    mime: str = "text/plain"
    metadata: dict[str, str] = Field(default_factory=dict)


def extract_docx(path: Path) -> ExtractedDocument:
    from docx import Document

    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    title = path.stem
    for para in doc.paragraphs:
        if para.style and para.style.name and "Heading" in para.style.name:
            if para.text.strip():
                title = para.text.strip()
                break
    return ExtractedDocument(
        source_type="docx",
        uri=str(path.resolve()),
        title=title,
        text=text or path.name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        metadata={"filename": path.name},
    )
