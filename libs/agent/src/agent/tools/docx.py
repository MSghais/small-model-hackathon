from __future__ import annotations

import re
import uuid
from pathlib import Path

from docx import Document
from docx.shared import Pt

from agent.models import SlideOutline
from agent.tools.pptx import _outputs_dir, _safe_filename


def create_docx(outline: SlideOutline, run_id: str | None = None) -> Path:
    """Build a Word document from a slide outline (opens in Google Docs when uploaded)."""
    rid = run_id or uuid.uuid4().hex[:12]
    out_dir = _outputs_dir()
    filename = f"{_safe_filename(outline.title)}_{rid}.docx"
    path = out_dir / filename

    doc = Document()
    title = doc.add_heading(outline.title, level=0)
    title.runs[0].font.size = Pt(28)
    doc.add_paragraph("Generated lesson slides").italic = True
    doc.add_page_break()

    for index, slide in enumerate(outline.slides, start=1):
        heading = doc.add_heading(f"Slide {index}: {slide.title}", level=1)
        heading.runs[0].font.size = Pt(22)
        for bullet in slide.bullets:
            para = doc.add_paragraph(bullet, style="List Bullet")
            para.runs[0].font.size = Pt(14)
        if slide.speaker_note:
            note = doc.add_paragraph(f"Teacher note: {slide.speaker_note}")
            note.runs[0].italic = True
            note.runs[0].font.size = Pt(11)
        if index < len(outline.slides):
            doc.add_page_break()

    doc.save(str(path))
    return path


def create_html_export(outline: SlideOutline, run_id: str | None = None) -> Path:
    """Standalone HTML file — import into Google Docs via File → Open → Upload."""
    from agent.preview import outline_to_html

    rid = run_id or uuid.uuid4().hex[:12]
    out_dir = _outputs_dir()
    filename = f"{_safe_filename(outline.title)}_{rid}.html"
    path = out_dir / filename

    body = outline_to_html(outline)
    full = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>{_escape_html(outline.title)}</title>
</head>
<body>
{body}
</body>
</html>
"""
    path.write_text(full)
    return path


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
