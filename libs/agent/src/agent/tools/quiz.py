"""Quiz export: DOCX worksheet + HTML preview."""

from __future__ import annotations

import html
from pathlib import Path

from docx import Document
from docx.shared import Pt

from agent.models import QuizOutline, QuizQuestion

_CHOICE_LABELS = ("A", "B", "C", "D")


def _add_question_docx(doc: Document, index: int, question: QuizQuestion) -> None:
    doc.add_paragraph(f"{index}. {question.prompt}")
    for label, choice in zip(_CHOICE_LABELS, question.choices, strict=True):
        doc.add_paragraph(f"   {label}. {choice}")
    doc.add_paragraph("")


def create_quiz_docx(outline: QuizOutline, path: Path) -> Path:
    """Student worksheet with numbered questions; answer key on final page."""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(outline.title, level=0)
    if outline.instructions.strip():
        doc.add_paragraph(outline.instructions.strip())
        doc.add_paragraph("")

    for i, question in enumerate(outline.questions, start=1):
        _add_question_docx(doc, i, question)

    doc.add_page_break()
    doc.add_heading("Answer Key", level=1)
    for i, question in enumerate(outline.questions, start=1):
        label = _CHOICE_LABELS[question.correct_index]
        answer = question.choices[question.correct_index]
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {label}. {answer}")
        run.bold = True
        if question.explanation.strip():
            doc.add_paragraph(question.explanation.strip(), style="List Bullet")

    doc.save(str(path))
    return path


def create_quiz_html(outline: QuizOutline, path: Path) -> Path:
    """Printable HTML worksheet with collapsible answer key."""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    title = html.escape(outline.title)
    instructions = html.escape(outline.instructions.strip()) if outline.instructions.strip() else ""

    question_blocks: list[str] = []
    answer_rows: list[str] = []

    for i, question in enumerate(outline.questions, start=1):
        prompt = html.escape(question.prompt)
        choices_html = "\n".join(
            f'<li><span class="choice-label">{label}.</span> {html.escape(choice)}</li>'
            for label, choice in zip(_CHOICE_LABELS, question.choices, strict=True)
        )
        question_blocks.append(
            f'<section class="question"><h3>{i}. {prompt}</h3><ol class="choices">{choices_html}</ol></section>'
        )
        correct_label = _CHOICE_LABELS[question.correct_index]
        correct_text = html.escape(question.choices[question.correct_index])
        expl = html.escape(question.explanation.strip()) if question.explanation.strip() else ""
        expl_html = f'<p class="explanation">{expl}</p>' if expl else ""
        answer_rows.append(
            f"<tr><td>{i}</td><td><strong>{correct_label}. {correct_text}</strong></td>"
            f"<td>{expl}</td></tr>"
            if expl
            else f"<tr><td>{i}</td><td><strong>{correct_label}. {correct_text}</strong></td><td></td></tr>"
        )

    body = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
  body {{ font-family: Georgia, serif; max-width: 720px; margin: 2rem auto; padding: 0 1rem; line-height: 1.5; }}
  h1 {{ font-size: 1.5rem; margin-bottom: 0.5rem; }}
  .instructions {{ margin-bottom 1.5rem; color: #333; }}
  .question {{ margin-bottom 1.25rem; page-break-inside: avoid; }}
  .question h3 {{ font-size: 1rem; font-weight: 600; margin: 0 0 0.5rem; }}
  .choices {{ list-style: none; padding-left: 0; margin: 0; }}
  .choices li {{ margin: 0.25rem 0; }}
  .choice-label {{ font-weight: 600; margin-right: 0.35rem; }}
  details.answer-key {{ margin-top: 2rem; border-top: 2px solid #ccc; padding-top: 1rem; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 0.9rem; }}
  th, td {{ border: 1px solid #ccc; padding: 0.4rem 0.6rem; text-align: left; vertical-align: top; }}
  th {{ background: #f5f5f5; }}
  @media print {{
    details.answer-key {{ display: block; }}
    details.answer-key summary {{ display: none; }}
  }}
</style>
</head>
<body>
  <h1>{title}</h1>
  {"<p class='instructions'>" + instructions + "</p>" if instructions else ""}
  {"".join(question_blocks)}
  <details class="answer-key">
    <summary>Answer key (click to expand)</summary>
    <table>
      <thead><tr><th>#</th><th>Answer</th><th>Explanation</th></tr></thead>
      <tbody>
        {"".join(answer_rows)}
      </tbody>
    </table>
  </details>
</body>
</html>
"""
    path.write_text(body, encoding="utf-8")
    return path


def create_quiz(outline: QuizOutline, output_dir: Path, stem: str = "quiz") -> dict[str, Path]:
    """Write DOCX and HTML exports; return paths keyed by format."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / f"{stem}.docx"
    html_path = output_dir / f"{stem}.html"
    create_quiz_docx(outline, docx_path)
    create_quiz_html(outline, html_path)
    return {"docx": docx_path, "html": html_path}
